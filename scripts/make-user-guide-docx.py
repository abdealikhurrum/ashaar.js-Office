#!/usr/bin/env python3
"""Generate docs/USER_GUIDE.docx from docs/USER_GUIDE.md (single source of truth).

Handles the Markdown subset used by the guide: # / ## / ### headings, - bullets,
1. ordered lists, ``` fenced code blocks, | pipe tables, and inline **bold** and
`code`. Arabic runs are marked RTL so they render correctly.

Run: python3 scripts/make-user-guide-docx.py  (or: npm run make-user-guide)
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "USER_GUIDE.md"
OUT = ROOT / "docs" / "USER_GUIDE.docx"

ARABIC = re.compile(r"[؀-ۿݐ-ݿﭐ-﷿ﹰ-﻿]")
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def is_arabic(text):
    return bool(ARABIC.search(text or ""))


def set_rtl_run(run):
    rpr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rpr.append(rtl)


def set_bidi(paragraph):
    ppr = paragraph._element.get_or_add_pPr()
    ppr.append(OxmlElement("w:bidi"))


def shade(paragraph, fill="F2F0EA"):
    ppr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_inline(paragraph, text, base_mono=False):
    """Add runs to a paragraph, handling **bold** and `code`; RTL for Arabic."""
    for seg in INLINE.split(text):
        if not seg:
            continue
        bold = mono = base_mono
        body = seg
        if seg.startswith("**") and seg.endswith("**"):
            bold, body = True, seg[2:-2]
        elif seg.startswith("`") and seg.endswith("`"):
            mono, body = True, seg[1:-1]
        run = paragraph.add_run(body)
        if bold:
            run.bold = True
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        if is_arabic(body):
            set_rtl_run(run)


def flush_code(doc, lines):
    if not lines:
        return
    p = doc.add_paragraph()
    shade(p)
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    arabic_block = any(is_arabic(l) for l in lines)
    if arabic_block:
        set_bidi(p)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        if is_arabic(line):
            set_rtl_run(run)


def flush_table(doc, rows):
    # rows: list of lists of cell strings; row[1] is the --- separator (skipped)
    header = rows[0]
    body = rows[2:] if len(rows) > 1 else []
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for j, cell in enumerate(header):
        para = table.rows[0].cells[j].paragraphs[0]
        add_inline(para, cell.strip(), base_mono=False)
        for r in para.runs:
            r.bold = True
    for r in body:
        cells = table.add_row().cells
        for j, cell in enumerate(r):
            if j < len(cells):
                add_inline(cells[j].paragraphs[0], cell.strip())
    doc.add_paragraph()


def parse_table_line(line):
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def main():
    md = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    in_code = False
    code_lines = []
    table_rows = []

    def flush_pending_table():
        nonlocal table_rows
        if table_rows:
            flush_table(doc, table_rows)
            table_rows = []

    for raw in md:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if in_code:
                flush_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_pending_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.lstrip().startswith("|"):
            table_rows.append(parse_table_line(line))
            continue
        else:
            flush_pending_table()

        if not line.strip():
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            heading = m.group(2)
            if level == 1:
                p = doc.add_heading("", level=0)
            else:
                p = doc.add_heading("", level=level - 1)
            add_inline(p, heading)
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            continue

        m = re.match(r"^[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1))
            continue

        p = doc.add_paragraph()
        add_inline(p, line)

    flush_pending_table()
    if in_code:
        flush_code(doc, code_lines)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
