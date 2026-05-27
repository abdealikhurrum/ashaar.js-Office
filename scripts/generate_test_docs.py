from __future__ import annotations

import json
from pathlib import Path
from zlib import crc32

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "test-documents"
USABLE_WIDTH = 9360
ACCENT = RGBColor(31, 111, 104)
REFRAIN = RGBColor(167, 53, 42)

POEM = """دل ناداں تجھے ہوا کیا ہے \\ آخر اس درد کی دوا کیا ہے
ہم ہیں مشتاق اور وہ بیزار \\ یا الہی یہ ماجرا کیا ہے

یا رب وہ نہ سمجھے ہیں نہ سمجھیں گے مری بات \\
دے اور دل ان کو جو نہ دے مجھ کو زباں اور
"""

SOLO_POEM = """دل ہی تو ہے نہ سنگ و خشت درد سے بھر نہ آئے کیوں

روئیں گے ہم ہزار بار کوئی ہمیں ستائے کیوں \\ %"""


def parse_poem(text: str) -> list[list[dict[str, str | bool | None]]]:
    stanzas = []
    for stanza in text.strip().split("\n\n"):
        rows = []
        pending = None
        for raw in stanza.splitlines():
            line = raw.strip()
            if not line:
                continue
            refrain = line.endswith("%")
            line = line.rstrip("%").strip()
            if line.endswith("\\"):
                pending = (line[:-1].strip(), refrain)
                continue
            if pending:
                rows.append({"sadr": pending[0], "ajuz": line, "sadr_refrain": pending[1], "ajuz_refrain": refrain})
                pending = None
            elif "\\" in line:
                sadr, ajuz = [part.strip() for part in line.split("\\", 1)]
                rows.append({"sadr": sadr, "ajuz": ajuz, "sadr_refrain": refrain, "ajuz_refrain": refrain})
            else:
                rows.append({"sadr": line, "ajuz": None, "sadr_refrain": refrain, "ajuz_refrain": False})
        if pending:
            rows.append({"sadr": pending[0], "ajuz": None, "sadr_refrain": pending[1], "ajuz_refrain": False})
        stanzas.append(rows)
    return stanzas


def visible_weight(text: str | None) -> float:
    text = (text or "").replace(" ", "").replace("ـ", "")
    return max(1, len(text))


def widths_for(stanza, layout: str, gap: int = 4) -> list[int]:
    if layout == "compact":
      return [750, 3745, 374, 3745, 746]
    if layout == "equal" or layout == "stacked":
        side = (USABLE_WIDTH - int(USABLE_WIDTH * gap / 100)) // 2
        return [side, USABLE_WIDTH - side * 2, side]

    max_sadr = max(visible_weight(row["sadr"]) for row in stanza if row["ajuz"])
    max_ajuz = max(visible_weight(row["ajuz"]) for row in stanza if row["ajuz"])
    gap_w = int(USABLE_WIDTH * gap / 100)
    available = USABLE_WIDTH - gap_w
    sadr = int(available * max_sadr / (max_sadr + max_ajuz))
    sadr = max(int(available * 0.36), min(int(available * 0.56), sadr))
    return [sadr, gap_w, available - sadr]


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    bidi = OxmlElement("w:bidiVisual")
    tbl_pr.append(bidi)

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells[: len(widths)]):
            set_cell_width(cell, widths[index])


def set_cell_margins(table, value=80):
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def rtl_paragraph(paragraph, alignment):
    paragraph.alignment = alignment
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.6


def write_cell(cell, text, alignment, refrain=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    rtl_paragraph(paragraph, alignment)
    run = paragraph.add_run(text or "")
    run.font.name = "Arial"
    run.font.size = Pt(15)
    if refrain:
        run.font.color.rgb = REFRAIN


def wrap_sdt(block_element, title, tag):
    parent = block_element.getparent()
    index = parent.index(block_element)
    parent.remove(block_element)

    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), title)
    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdt_pr.append(alias)
    sdt_pr.append(tag_el)
    sdt_content = OxmlElement("w:sdtContent")
    sdt_content.append(block_element)
    sdt.append(sdt_pr)
    sdt.append(sdt_content)
    parent.insert(index, sdt)


def add_poem_table(doc, stanza, layout, source):
    widths = widths_for(stanza, layout)
    table = doc.add_table(rows=0, cols=len(widths))
    table.autofit = False
    set_table_geometry(table, widths)
    set_cell_margins(table)

    for bayt in stanza:
        ajuz = bayt["ajuz"]
        if layout == "stacked" or not ajuz:
            row = table.add_row()
            merged = row.cells[0]
            for cell in row.cells[1:]:
                merged = merged.merge(cell)
            text = str(bayt["sadr"])
            if ajuz:
                text += "\n" + str(ajuz)
            write_cell(merged, text, WD_ALIGN_PARAGRAPH.CENTER, bool(bayt["sadr_refrain"]))
            continue

        row = table.add_row()
        offset = 1 if layout == "compact" else 0
        if layout == "compact":
            write_cell(row.cells[0], "", WD_ALIGN_PARAGRAPH.CENTER)
            write_cell(row.cells[4], "", WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[offset], str(bayt["sadr"]), WD_ALIGN_PARAGRAPH.LEFT, bool(bayt["sadr_refrain"]))
        write_cell(row.cells[offset + 1], "", WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[offset + 2], str(ajuz), WD_ALIGN_PARAGRAPH.RIGHT, bool(bayt["ajuz_refrain"]))

    payload = {
        "k": "ashaar-poem",
        "v": 1,
        "layoutMode": layout,
        "widthMode": "optimized" if layout == "balanced" else "fixed",
        "sourceHash": f"{crc32(source.encode('utf-8')) & 0xffffffff:x}",
    }
    wrap_sdt(table._tbl, "Ashaar Poem", "ashaar:" + json.dumps(payload, separators=(",", ":")))
    doc.add_paragraph()


def base_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(subtitle)
    p.paragraph_format.space_after = Pt(18)
    return doc


def build_doc(filename, layout, text):
    doc = base_doc(f"Ashaar {layout.title()} Table Test", "Generated fixture with Word tables wrapped in Ashaar content controls.")
    for stanza in parse_poem(text):
        add_poem_table(doc, stanza, layout, text)
    doc.save(OUT / filename)


def main():
    OUT.mkdir(exist_ok=True)
    build_doc("ashaar-balanced-table.docx", "balanced", POEM)
    build_doc("ashaar-equal-table.docx", "equal", POEM)
    build_doc("ashaar-compact-table.docx", "compact", POEM)
    build_doc("ashaar-stacked-content-control.docx", "stacked", SOLO_POEM)


if __name__ == "__main__":
    main()
